import os
import re
import io
import datetime
import pandas as pd
import plotly.express as px
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import stripe
import chromadb
from supabase import create_client, Client

# --- 1. STREAMLIT PAGE SETUP (MUST BE FIRST) ---
st.set_page_config(page_title="QuantLex", page_icon="📊", layout="wide")

# --- 2. WAKE UP API KEYS ---
supabase_url = st.secrets["SUPABASE_URL"]
supabase_key = st.secrets["SUPABASE_KEY"]
supabase: Client = create_client(supabase_url, supabase_key)
stripe.api_key = st.secrets["STRIPE_API_KEY"]

# --- 3. PREVENT DATABASE CRASHES ---
@st.cache_resource
def get_chroma_client():
    import chromadb.api
    chromadb.api.client.SharedSystemClient.clear_system_cache()
    return chromadb.PersistentClient(path="./chroma_db")

chroma_client = get_chroma_client()
collection = chroma_client.get_or_create_collection(name="financial_vault")

# --- 4. THE STRIPE AUTO-UNLOCKER ---
if "session_id" in st.query_params:
    session_id = st.query_params["session_id"]
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        if session.payment_status == "paid" and session.client_reference_id:
            user_id = session.client_reference_id
            
            # Unlock database
            supabase.rpc("unlock_user_subscription", {"target_user_id": user_id}).execute()
            
            # VIP Auto-Login (forces them past the login screen on the new tab)
            class PaidUser:
                def __init__(self, uid):
                    self.id = uid
                    self.email = "Upgraded User" # Placeholder for auto-login
            st.session_state["user"] = PaidUser(user_id)
            
            # Clear URL so it doesn't run twice
            st.query_params.clear()
            st.rerun()
    except Exception:
        pass

# --- 5. THE SINGLE LOGIN SCREEN ---
if "user" not in st.session_state:
    st.title("Enterprise RAG Access")
    tab1, tab2 = st.tabs(["Login", "Create Account"])
    
    with tab1:
        log_email = st.text_input("Email", key="log_email")
        log_pwd = st.text_input("Password", type="password", key="log_pwd")
        if st.button("Login"):
            try:
                res = supabase.auth.sign_in_with_password({"email": log_email, "password": log_pwd})
                st.session_state["user"] = res.user
                st.rerun()
            except Exception as e:
                st.error("Login failed: Invalid email or password.")
                
    with tab2:
        reg_email = st.text_input("Email", key="reg_email")
        reg_pwd = st.text_input("Password", type="password", key="reg_pwd")
        if st.button("Sign Up"):
            try:
                res = supabase.auth.sign_up({"email": reg_email, "password": reg_pwd})
                st.success("Account created successfully. Please log in.")
            except Exception as e:
                st.error(f"Registration failed: {e}")
                
    st.stop() # Stops the app here if they aren't logged in

# --- 6. THE STRIPE PAYWALL ---
user_id = st.session_state["user"].id
profile = supabase.table("profiles").select("is_subscribed").eq("id", user_id).execute()

if len(profile.data) == 0:
    is_subscribed = False
else:
    is_subscribed = profile.data[0].get("is_subscribed", False)

if not is_subscribed:
    st.title("🔒 Enterprise Plan Required")
    st.warning("Your account is currently inactive. Upgrade to unlock the secure RAG vault.")
    
    st.link_button(
        "Upgrade to Enterprise ($499/mo)", 
        f"{st.secrets['STRIPE_PAYMENT_LINK']}?client_reference_id={user_id}", 
        type="primary", 
        use_container_width=True
    )
    
    if st.button("I just paid - Refresh my account", type="secondary"):
        st.rerun()
        
    st.stop() # Stops them from seeing the AI Chat until paid

# =================================================================
# --- RAG UI CHAT INTERFACE ---
# =================================================================
st.title("QuantLex")
st.caption("Powered by OpenAI GPT-4o, ChromaDB, Pandas, Supabase Vaults & Word Brief Generation")

# Setup user variables for the rest of the app based on Supabase
current_user = str(st.session_state["user"].id)
name = getattr(st.session_state["user"], "email", "User")

# --- SECURE SIDEBAR UPLOADER ---
with st.sidebar:
    st.header("📂 Secure Workspace")
    st.write(f"Logged in as: {name}")
    
    uploaded_file = st.file_uploader("Upload a text document", type=["txt"])
    
    if uploaded_file is not None and st.button("Vault Document"):
        with st.spinner("Encrypting and vaulting..."):
            text_data = uploaded_file.getvalue().decode("utf-8")
            filename = uploaded_file.name
            chunk_size = 1000
            chunks = [text_data[i:i+chunk_size] for i in range(0, len(text_data), chunk_size)]
            
            for i, text_chunk in enumerate(chunks):
                chunk_id = f"{current_user}_{filename}_chunk_{i}"
                collection.add(
                    documents=[text_chunk],
                    metadatas=[{"user_id": current_user, "source": filename}],
                    ids=[chunk_id]
                )
            st.success(f"Successfully vaulted: {filename}")

# =====================================================================
# SUPABASE CLOUD VAULT ENGINE
# =====================================================================
BUCKET_NAME = "client-vaults"

def upload_file_to_supabase(user_id, local_file_path, filename):
    cloud_path = f"{user_id}/{filename}"
    with open(local_file_path, "rb") as f:
        file_bytes = f.read()
    try:
        supabase.storage.from_(BUCKET_NAME).upload(
            file=file_bytes,
            path=cloud_path,
            file_options={"upsert": "true"}
        )
    except Exception as e:
        print(f"Supabase Upload Exception for {filename}: {e}")

def delete_file_from_supabase(user_id, filename):
    cloud_path = f"{user_id}/{filename}"
    try:
        res = supabase.storage.from_(BUCKET_NAME).remove([cloud_path])
        if not res:
            st.warning(f"⚠️ Supabase could not delete '{cloud_path}'. Verify permissions.")
        else:
            st.toast(f"☁️ Cloud Vault synced: Deleted '{filename}'")
        return res
    except Exception as e:
        st.error(f"❌ Cloud Deletion Blocked: {e}")
        return None

def sync_supabase_to_local(user_id, local_docs_folder):
    try:
        remote_files = supabase.storage.from_(BUCKET_NAME).list(user_id)
        for item in remote_files:
            filename = item.get("name")
            if not filename or filename == ".emptyFolderPlaceholder":
                continue
            
            local_path = os.path.join(local_docs_folder, filename)
            if not os.path.exists(local_path):
                cloud_path = f"{user_id}/{filename}"
                file_bytes = supabase.storage.from_(BUCKET_NAME).download(cloud_path)
                with open(local_path, "wb") as f:
                    f.write(file_bytes)
    except Exception as e:
        print(f"Supabase Sync Warning: {e}")

# =====================================================================
# HYBRID ENGINE SETUP 
# =====================================================================
@st.cache_resource(show_spinner=False)
def load_rag_engine(user_id):
    engine_dir = os.path.dirname(os.path.abspath(__file__))
    db_path = os.path.join(engine_dir, "rag_db", "tenants", user_id)
    docs_folder = os.path.join(engine_dir, "rag documents", "tenants", user_id)

    if not os.path.exists(docs_folder):
        os.makedirs(docs_folder)
    if not os.path.exists(db_path):
        os.makedirs(db_path)

    sync_supabase_to_local(user_id, docs_folder)
    chroma_client = chromadb.PersistentClient(path=db_path)
    collection = chroma_client.get_or_create_collection(name=f"collection_{user_id}")
    openai_client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
    
    def chunk_text(text, chunk_size=150):
        words = text.split()
        return [" ".join(words[i : i + chunk_size]) for i in range(0, len(words), chunk_size)]

    chunks_loaded = 0
    spreadsheet_tables = []

    if os.path.exists(docs_folder):
        all_chunks = []
        all_ids = []
        chunk_counter = 0

        for filename in os.listdir(docs_folder):
            file_path = os.path.join(docs_folder, filename)
            text_content = ""

            if filename.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as file:
                    text_content = file.read()

            elif filename.endswith(".pdf"):
                try:
                    reader = PdfReader(file_path)
                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text_content += extracted + " "
                except Exception as e:
                    print(f"Error reading PDF {filename}: {e}")

            elif filename.endswith(".csv") or filename.endswith(".xlsx"):
                try:
                    if filename.endswith(".csv"):
                        df = pd.read_csv(file_path)
                    else:
                        df = pd.read_excel(file_path)
                    df = df.dropna(how="all").fillna("N/A")
                    table_string = f"=== SPREADSHEET DATA FROM [{filename}] ===\n{df.to_string(index=False)}"
                    spreadsheet_tables.append(table_string)
                except Exception as e:
                    print(f"Error reading spreadsheet {filename}: {e}")

            if text_content:
                chunks = chunk_text(text_content, chunk_size=150)
                for chunk in chunks:
                    all_chunks.append(chunk)
                    all_ids.append(f"{filename}_chunk_{chunk_counter}")
                    chunk_counter += 1

        if all_chunks:
            collection.upsert(documents=all_chunks, ids=all_ids)
            chunks_loaded = len(all_chunks)

    return collection, openai_client, chunks_loaded, spreadsheet_tables, docs_folder

tenant_collection, openai_client, chunks_loaded, spreadsheet_tables, docs_folder = load_rag_engine(current_user)

# =====================================================================
# THE SIDEBAR DASHBOARD
# =====================================================================
with st.sidebar:
    st.divider()
    if st.button("🚪 Log Out", use_container_width=True):
        st.session_state.clear()
        st.rerun()
    st.markdown("---")

    st.header("📊 System Dashboard")
    st.success("🟢 Cloud Vault & Brief Generator Online")

    st.markdown("---")
    st.subheader("Database Stats")
    st.write(f"**Text Paragraphs Loaded:** `{chunks_loaded}`")
    st.write(f"**Active Spreadsheets Loaded:** `{len(spreadsheet_tables)}`")

    st.markdown("---")
    st.subheader("📥 Add New Knowledge")

    uploaded_files = st.file_uploader(
        "Upload TXT, PDF, CSV, or XLSX files:",
        type=["txt", "pdf", "csv", "xlsx"],
        accept_multiple_files=True,
    )

    if uploaded_files and st.button("🚀 Process & Ingest Files", use_container_width=True):
        with st.spinner("Vaulting files to Supabase & Hybrid Engines..."):
            for uploaded_file in uploaded_files:
                file_path = os.path.join(docs_folder, uploaded_file.name)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                upload_file_to_supabase(current_user, file_path, uploaded_file.name)

            load_rag_engine.clear()
            st.success("Files successfully processed and vaulted!")
            st.rerun()

    st.markdown("---")
    st.subheader("🗂️ Manage Uploaded Docs")
    existing_files = os.listdir(docs_folder) if os.path.exists(docs_folder) else []

    if existing_files:
        file_to_delete = st.selectbox("Select a document to remove:", existing_files)
        if st.button("🗑️ Delete Selected File", use_container_width=True):
            with st.spinner("Removing file & purging database memory..."):
                file_path = os.path.join(docs_folder, file_to_delete)
                if os.path.exists(file_path):
                    os.remove(file_path)

                delete_file_from_supabase(current_user, file_to_delete)

                db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rag_db", "tenants", current_user)
                chroma_client = chromadb.PersistentClient(path=db_path)
                try:
                    chroma_client.delete_collection(name=f"collection_{current_user}")
                except Exception:
                    pass

                load_rag_engine.clear()
                st.success(f"Successfully deleted '{file_to_delete}'!")
                st.rerun()
    else:
        st.info("No documents uploaded yet.")

    st.markdown("---")
    if st.button("🗑️ Clear Chat History", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

    if "messages" in st.session_state and st.session_state.messages:
        chat_transcript = "=== ENTERPRISE AI CHAT LOG ===\n\n"
        for msg in st.session_state.messages:
            role = "👤 You" if msg["role"] == "user" else "🤖 AI Assistant"
            chat_transcript += f"{role}:\n{msg['content']}\n\n"

        st.download_button(
            label="💾 Download Chat Transcript (.txt)",
            data=chat_transcript,
            file_name=f"chat_log_{current_user}.txt",
            mime="text/plain",
            use_container_width=True,
        )

# =====================================================================
# INTERACTIVE CHART RENDERING & CLEANING ENGINE
# =====================================================================
def clean_ai_text(text):
    if not text:
        return ""
    return re.sub(r"```(?:python)?\s*.*?```", "", text, flags=re.DOTALL | re.IGNORECASE).strip()

@st.cache_data(show_spinner=False)
def create_executive_brief_docx(query_text, ai_response, user_name):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    title = doc.add_paragraph()
    title_run = title.add_run("EXECUTIVE BRIEFING & AUDIT REPORT")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_paragraph()
    meta.add_run("Prepared For: ").bold = True
    meta.add_run(f"{user_name}\n")
    meta.add_run("Inquiry / Subject: ").bold = True
    meta.add_run(f"{query_text}\n")
    
    clean_text = clean_ai_text(ai_response)
    doc.add_paragraph("_________________________________________________________________________________")
    
    for line in clean_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.add_run(line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def render_ai_charts(text, msg_idx=0):
    code_blocks = re.findall(r"```(?:python)?\s*(.*?)\s*```", text, re.DOTALL | re.IGNORECASE)
    for i, code in enumerate(code_blocks):
        if "fig" in code:
            try:
                clean_code = re.sub(r"fig\.show\(\)", "", code)
                exec_scope = {"pd": pd, "px": px, "st": st}
                exec(clean_code, exec_scope)
                fig = exec_scope.get("fig") or exec_scope.get("figure")
                if fig:
                    st.plotly_chart(fig, use_container_width=True)
            except Exception as e:
                st.caption(f"⚠️ *Chart rendering note: {e}*")

def live_stream_filter(stream_obj, raw_accumulator):
    hide_code = False
    buffer = ""
    for chunk in stream_obj:
        if not hasattr(chunk, "choices") or not chunk.choices:
            continue
        delta = getattr(chunk.choices[0], "delta", None)
        if not delta:
            continue
        token = getattr(delta, "content", "") or ""
        if not token:
            continue
        raw_accumulator.append(token)
        if not hide_code:
            buffer += token
            if "```" in buffer:
                hide_code = True
                clean_prefix = buffer.split("```")[0]
                if clean_prefix:
                    yield clean_prefix
                buffer = ""
            elif not buffer.endswith("`"):
                yield buffer
                buffer = ""
    if not hide_code and buffer:
        yield buffer

# =====================================================================
# CONVERSATIONAL MEMORY
# =====================================================================
if "messages" not in st.session_state:
    st.session_state.messages = []

for idx, message in enumerate(st.session_state.messages):
    with st.chat_message(message["role"]):
        if message["role"] == "assistant":
            st.markdown(clean_ai_text(message["content"]))
            render_ai_charts(message["content"], msg_idx=idx)
            
            associated_query = "General Executive Inquiry"
            if idx > 0 and st.session_state.messages[idx - 1]["role"] == "user":
                associated_query = st.session_state.messages[idx - 1]["content"]
                
            docx_bytes = create_executive_brief_docx(associated_query, message["content"], name)
            st.download_button(
                label="📄 Download Executive Brief (.docx)",
                data=docx_bytes,
                file_name=f"Executive_Brief_{idx}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"brief_btn_hist_{idx}",
                use_container_width=False,
            )
        else:
            st.markdown(message["content"])

# =====================================================================
# THE INTERACTIVE CHAT BOX 
# =====================================================================
if query := st.chat_input("Ask about policies, syllabi, or generate spreadsheet charts..."):

    with st.chat_message("user"):
        st.markdown(query)
    st.session_state.messages.append({"role": "user", "content": query})

    results = tenant_collection.query(
        query_texts=[query], 
        n_results=15
    )
    
    text_context = "\n---\n".join(results["documents"][0]) if results["documents"] and results["documents"][0] else "No text documents found."
    spreadsheet_context = "\n\n".join(spreadsheet_tables) if spreadsheet_tables else "No spreadsheet tables loaded."
    full_combined_context = f"--- UNSTRUCTURED TEXT DOCUMENTS ---\n{text_context}\n\n--- STRUCTURED SPREADSHEETS (100% INTACT TABLES) ---\n{spreadsheet_context}"

    messages_payload = [{"role": "system", "content": "You are an elite enterprise AI data assistant."}]
    
    for msg in st.session_state.messages[-6:-1]:
        messages_payload.append({"role": msg["role"], "content": msg["content"]})

    messages_payload.append(
        {"role": "user", "content": f"System Data Context:\n{full_combined_context}\n\nCurrent Question: {query}"}
    )

    with st.chat_message("assistant"):
        try:
            raw_tokens = []
            stream = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=messages_payload,
                stream=True,
                temperature=0.0,
            )

            st.write_stream(live_stream_filter(stream, raw_tokens))
            full_ai_answer = "".join(raw_tokens)

            st.session_state.messages.append({"role": "assistant", "content": full_ai_answer})
            st.rerun()

        except Exception as e:
            st.error(f"Error connecting to AI: {e}")